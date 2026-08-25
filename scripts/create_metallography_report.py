from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


THRESHOLDS = {"iou": 0.85, "precision": 0.90, "recall": 0.95}
OUTPUT_NAME = "AI金相检测项目性能验证报告.docx"
METRICS = ("iou", "precision", "recall")
METRIC_LABELS = {"iou": "IoU", "precision": "P", "recall": "R"}
COLORS = {"iou": "#2E74B5", "precision": "#3A8F5C", "recall": "#D9822B"}

_FONT_PATHS = (
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
)
_FONT_PATH = next((p for p in _FONT_PATHS if p.exists()), None)
_FONT_NAME = font_manager.FontProperties(fname=str(_FONT_PATH)).get_name() if _FONT_PATH else "DejaVu Sans"


def build_summary() -> pd.DataFrame:
    """Return the seven customer-provided validation results used in the report."""

    rows = [
        ("夹杂物", "条状", 1570, 0.871, 0.902, 0.959),
        ("夹杂物", "点状", 2311, 0.895, 0.929, 0.968),
        ("晶粒度", "晶界", 401, 0.951, 0.960, 0.979),
        ("脱碳层", "镶嵌-全脱", 56, 0.920, 0.935, 0.967),
        ("脱碳层", "镶嵌-总脱", 217, 0.912, 0.906, 0.953),
        ("脱碳层", "非镶嵌-全脱", 61, 0.915, 0.925, 0.958),
        ("脱碳层", "非镶嵌-总脱", 398, 0.941, 0.961, 0.970),
    ]
    frame = pd.DataFrame(rows, columns=["project", "label", "sample_count", *METRICS])
    frame["pass"] = frame["iou"].ge(THRESHOLDS["iou"]) & frame["precision"].ge(
        THRESHOLDS["precision"]
    ) & frame["recall"].ge(THRESHOLDS["recall"])
    return frame


def _font_kwargs() -> dict[str, object]:
    return {"fontproperties": font_manager.FontProperties(fname=str(_FONT_PATH))} if _FONT_PATH else {}


def _configure_matplotlib() -> None:
    plt.rcParams["font.family"] = _FONT_NAME
    plt.rcParams["axes.unicode_minus"] = False


def _plot_grouped_bars(frame: pd.DataFrame, title: str, path: Path) -> None:
    _configure_matplotlib()
    fig, ax = plt.subplots(figsize=(10, 5.2), dpi=180)
    x = list(range(len(frame)))
    width = 0.23
    for offset, metric in enumerate(METRICS):
        values = frame[metric].tolist()
        bars = ax.bar(
            [value + (offset - 1) * width for value in x],
            values,
            width=width,
            color=COLORS[metric],
            label=METRIC_LABELS[metric],
        )
        for bar, value in zip(bars, values):
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                value + 0.004,
                f"{value:.3f}",
                ha="center",
                va="bottom",
                fontsize=8,
                **_font_kwargs(),
            )
    for metric in METRICS:
        ax.axhline(
            THRESHOLDS[metric],
            color=COLORS[metric],
            linestyle="--",
            linewidth=0.8,
            alpha=0.5,
        )
    labels = [str(label).replace("-", "\n") for label in frame["label"]]
    ax.set_xticks(x, labels, **_font_kwargs())
    ax.set_ylim(0.80, 1.02)
    ax.set_ylabel("指标值", **_font_kwargs())
    ax.set_title(title, fontsize=14, fontweight="bold", **_font_kwargs())
    ax.grid(axis="y", color="#D9E2EC", linewidth=0.7)
    ax.set_axisbelow(True)
    legend_kwargs = {"prop": font_manager.FontProperties(fname=str(_FONT_PATH))} if _FONT_PATH else {}
    ax.legend(ncol=3, loc="upper center", bbox_to_anchor=(0.5, -0.12), frameon=False, **legend_kwargs)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_charts(summary: pd.DataFrame, chart_dir: Path) -> list[Path]:
    chart_dir.mkdir(parents=True, exist_ok=True)
    paths = [
        chart_dir / "overview_metrics.png",
        chart_dir / "inclusion_metrics.png",
        chart_dir / "grain_boundary_metrics.png",
        chart_dir / "decarb_metrics.png",
    ]
    _plot_grouped_bars(summary, "三个检测项目性能指标总览", paths[0])
    _plot_grouped_bars(summary.iloc[:2], "夹杂物检测性能指标", paths[1])
    _plot_grouped_bars(summary.iloc[2:3], "晶粒度检测性能指标", paths[2])
    _plot_grouped_bars(summary.iloc[3:], "脱碳层检测性能指标", paths[3])
    return paths


def _set_run_font(run, name: str = "宋体", size: float | None = None, bold: bool | None = None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_widths(table, widths: Iterable[float]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_metric_table(doc: Document, frame: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["验证单元", "测试样本数", "IoU", "P", "R", "是否达到阈值"]
    for cell, text in zip(table.rows[0].cells, headers):
        _set_cell_shading(cell, "D9E2F3")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _set_run_font(run, size=10, bold=True)
    for project, label, sample_count, iou, precision, recall, passed in frame.itertuples(index=False, name=None):
        cells = table.add_row().cells
        values = [label, str(sample_count), f"{iou:.3f}", f"{precision:.3f}", f"{recall:.3f}", "是" if passed else "否"]
        for cell, text in zip(cells, values):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            _set_run_font(run, size=10)
    _set_table_widths(table, (2.05, 1.05, 0.75, 0.75, 0.75, 1.15))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_chart(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.1))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    _set_run_font(caption_run, size=9, color=(89, 89, 89))


def _add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def _setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(24)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Heading 1", 16, (46, 116, 181), 16, 8),
        ("Heading 2", 13, (46, 116, 181), 12, 6),
        ("Heading 3", 12, (31, 77, 120), 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("AI 金相检测系统｜性能验证报告")
    _set_run_font(header_run, size=9, color=(100, 100, 100))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("第 ")
    _set_run_font(footer_run, size=9, color=(100, 100, 100))
    _add_page_field(footer)
    footer_run = footer.add_run(" 页")
    _set_run_font(footer_run, size=9, color=(100, 100, 100))
    return doc


def _add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        _set_run_font(first, size=10.5, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        _set_run_font(rest, size=10.5)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, size=10.5)


def _add_section(doc: Document, heading: str, frame: pd.DataFrame, chart_path: Path, chart_caption: str, analysis: str) -> None:
    doc.add_heading(heading, level=1)
    _add_paragraph(doc, "本节根据客户填写的真实性能指标数据，展示该检测项目的 IoU、Precision（P）和 Recall（R）结果。")
    _add_metric_table(doc, frame)
    _add_chart(doc, chart_path, chart_caption)
    _add_paragraph(doc, analysis)


def build_report(output_dir: Path) -> Path:
    output_dir = Path(output_dir)
    chart_dir = output_dir / "charts"
    output_dir.mkdir(parents=True, exist_ok=True)
    summary = build_summary()
    summary.to_csv(output_dir / "metallography_summary.csv", index=False, encoding="utf-8-sig")
    chart_paths = create_charts(summary, chart_dir)

    doc = _setup_document()
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(90)
    title = cover.add_run("AI 金相检测项目\n性能验证报告")
    _set_run_font(title, name="黑体", size=25, bold=True, color=(31, 77, 120))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(16)
    subtitle_run = subtitle.add_run("夹杂物｜晶粒度｜脱碳层")
    _set_run_font(subtitle_run, size=14, color=(80, 80, 80))
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.paragraph_format.space_before = Pt(100)
    date_run = date_paragraph.add_run("2026 年 8 月")
    _set_run_font(date_run, size=11, color=(100, 100, 100))
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(16)
    notice_run = notice.add_run("数据来源：客户填写的真实性能指标数据。")
    _set_run_font(notice_run, size=10, bold=True, color=(31, 77, 120))
    doc.add_page_break()

    doc.add_heading("项目概述", level=1)
    _add_paragraph(doc, "本报告面向客户验收与性能验证，集中展示三个 AI 金相检测项目的核心评价指标。")
    _add_paragraph(doc, "夹杂物检测包含条状和点状两组；晶粒度检测包含晶界一组；脱碳层检测包含镶嵌-全脱、镶嵌-总脱、非镶嵌-全脱和非镶嵌-总脱四组。")

    doc.add_heading("验证方法", level=1)
    _add_paragraph(doc, "本报告采用标准中常用的三项指标进行展示：IoU 表示预测区域与参考区域的交并比；P（Precision）表示预测为目标区域的结果中实际正确的比例；R（Recall）表示参考目标区域被正确识别的比例。")
    _add_paragraph(doc, "本报告采用的展示阈值为：IoU ≥ 0.85，P ≥ 0.90，R ≥ 0.95。表格中的“是”表示该验证单元的三项指标均达到对应阈值。")

    doc.add_heading("执行摘要", level=1)
    _add_paragraph(doc, "下表为 7 个验证单元的 IoU、P、R 实测汇总结果。")
    _add_metric_table(doc, summary)
    _add_chart(doc, chart_paths[0], "图 1  三个检测项目核心指标总览")

    _add_section(
        doc,
        "夹杂物性能验证",
        summary.iloc[:2],
        chart_paths[1],
        "图 2  夹杂物条状/点状指标对比",
        "条状夹杂物的 IoU（0.871）和 P（0.902）低于点状夹杂物（IoU 0.895、P 0.929），主要原因是条状目标长度、宽度和走向变化更大，边界容易出现断裂、粘连或局部偏移；R 仍达到 0.959，说明目标主体检出较充分，但边界外扩和误检使 P 与 IoU 受到影响。点状目标形态相对集中，定位和分割更稳定，因此三项指标整体更高。",
    )
    _add_section(
        doc,
        "晶粒度性能验证",
        summary.iloc[2:3],
        chart_paths[2],
        "图 3  晶界指标结果",
        "晶界的 IoU、P、R 分别为 0.951、0.960 和 0.979，为本次各项任务中整体最高的结果。晶界在图像中通常具有连续的网络结构和较稳定的纹理/对比度，连续性信息有利于抑制孤立误检；同时本组测试样本数为 401，标注和样本覆盖相对充分，因此边界重合度和召回率均较高。",
    )
    _add_section(
        doc,
        "脱碳层性能验证",
        summary.iloc[3:],
        chart_paths[3],
        "图 4  脱碳层四组工况指标对比",
        "镶嵌全脱的 IoU/P/R 为 0.920/0.935/0.967，高于镶嵌总脱的 0.912/0.906/0.953，说明全脱区域组织特征更明确，而总脱终点处于渐变过渡带，判定边界更容易产生偏差。非镶嵌全脱的 P 为 0.925，低于镶嵌全脱，符合异形边缘、倒角、氧化层和照明变化带来的额外误检风险。值得注意的是，本批非镶嵌总脱的 IoU/P/R 达到 0.941/0.961/0.970，反而高于非镶嵌全脱；结合其 398 个测试样本，说明本批总脱样本的边界特征和标注一致性较好，不能仅凭任务名称预设其指标一定更低。",
    )

    doc.add_heading("总体验收结论", level=1)
    _add_paragraph(doc, "根据本次填写的真实性能指标数据，夹杂物、晶粒度和脱碳层三个检测项目的 7 个验证单元，其 IoU、P、R 三项指标均达到设定阈值。")
    _add_paragraph(doc, "从整体指标关系看，R 普遍高于 P，说明系统对目标的覆盖能力较好，但仍存在一定程度的误检、边界外扩或区域归属偏差；后续优化重点应放在条状夹杂物边界精修、总脱过渡带终点判定及非镶嵌样品边缘伪影抑制。")

    doc.add_heading("指标定义", level=1)
    _add_paragraph(doc, "IoU = 预测区域与参考区域的交集面积 ÷ 预测区域与参考区域的并集面积。")
    _add_paragraph(doc, "P = TP ÷ (TP + FP)；R = TP ÷ (TP + FN)。")

    output_path = output_dir / OUTPUT_NAME
    doc.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate the concise AI metallography validation report.")
    parser.add_argument("--output-dir", type=Path, default=Path("output"))
    args = parser.parse_args()
    print(build_report(args.output_dir))


if __name__ == "__main__":
    main()
