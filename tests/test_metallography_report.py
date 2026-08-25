from pathlib import Path

import pandas as pd
from docx import Document

from scripts.create_metallography_report import build_report, build_summary, create_charts


def test_summary_has_seven_units_and_three_metrics():
    summary = build_summary()
    assert len(summary) == 7
    assert {"project", "label", "iou", "precision", "recall"} <= set(summary.columns)
    assert (summary["iou"] >= 0.85).all()
    assert (summary["precision"] >= 0.90).all()
    assert (summary["recall"] >= 0.95).all()
    decarb = summary[summary["project"] == "脱碳层"].set_index("label")
    assert decarb.loc["镶嵌-总脱", "iou"] == 0.941
    assert decarb.loc["非镶嵌-总脱", "iou"] == 0.912


def test_charts_are_created(tmp_path: Path):
    paths = create_charts(build_summary(), tmp_path)
    assert {p.name for p in paths} == {
        "overview_metrics.png",
        "inclusion_metrics.png",
        "grain_boundary_metrics.png",
        "decarb_metrics.png",
    }
    assert all(p.stat().st_size > 10_000 for p in paths)


def test_docx_contains_required_sections(tmp_path: Path):
    output = build_report(tmp_path)
    doc = Document(output)
    text = "\n".join(p.text for p in doc.paragraphs)
    for heading in (
        "执行摘要",
        "验证方法",
        "夹杂物性能验证",
        "晶粒度性能验证",
        "脱碳层性能验证",
        "总体验收结论",
    ):
        assert heading in text
    assert len(doc.tables) >= 4
    assert len(doc.inline_shapes) == 4


def test_report_summary_matches_fixed_data(tmp_path: Path):
    output_dir = tmp_path / "report"
    build_report(output_dir)
    summary = pd.read_csv(output_dir / "metallography_summary.csv")
    assert len(summary) == 7
    assert summary["pass"].eq(True).all()
    assert (output_dir / "AI金相检测项目性能验证报告.docx").stat().st_size > 50_000
